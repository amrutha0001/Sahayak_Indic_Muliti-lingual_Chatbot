"""
ingest.py
---------
Reads scheme files from /schemes folder AND optionally crawls URLs.

Supported input formats:
  .txt  — plain text
  .pdf  — via pdfplumber
  .docx — via python-docx
  .doc  — via python-docx (best effort)
  .md   — markdown plain text
  .html — HTML files (via BeautifulSoup)
  .rtf  — RTF plain text extraction

URL crawling:
  Place a file called `urls.txt` inside /schemes with one URL per line.
  The script will fetch each URL, strip HTML, follow anchor links,
  and ingest the content (depth-1 crawl).

Run once (or whenever you add new files / URLs):
    python ingest.py
"""

import os
import re
import sys
import chromadb
import requests
from chromadb.utils import embedding_functions
from dotenv import load_dotenv
from urllib.parse import urljoin, urlparse

load_dotenv()

SCHEMES_DIR = "./schemes"
CHROMA_DIR  = "./chroma_db"
EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

# ── Optional imports ──────────────────────────────────────────────────────────
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    print("⚠️  pdfplumber not found — PDF ingestion disabled. pip install pdfplumber")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️  python-docx not found — DOCX ingestion disabled. pip install python-docx")

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False
    print("⚠️  beautifulsoup4 not found — URL/HTML crawling disabled. pip install beautifulsoup4")

# ── Embedding + ChromaDB ──────────────────────────────────────────────────────
sentence_ef = embedding_functions.SentenceTransformerEmbeddingFunction(
    model_name=EMBED_MODEL
)
client = chromadb.PersistentClient(path=CHROMA_DIR)
try:
    client.delete_collection("schemes")
except:
    pass
collection = client.create_collection(
    name="schemes",
    embedding_function=sentence_ef,
    metadata={"hnsw:space": "cosine"}
)


# ═══════════════════════════════════════════════════════════════════════════════
#  TEXT EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════

def extract_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _clean_pdf_text(text: str) -> str:
    """
    Fix common PDF extraction artefacts:
      - (cid:NNN) glyphs → Unicode (bullets, ligatures)
      - Font-mapped ₹ variants: ■, □, ▪, or bare 'n' before amounts → ₹
    """
    import re as _r
    cid_map = {
        127: "•", 183: "•", 122: "•", 108: "fi", 109: "fl",
        61623: "•", 61664: "•", 176: "°", 169: "©", 174: "®",
    }
    def _rep(m):
        return cid_map.get(int(m.group(1)), "•")

    text = _r.sub(r"\(cid:(\d+)\)", _rep, text)
    # Block elements used as ₹ in some PDF fonts
    text = _r.sub(r"[■□▪▸](?=\s*[\d,])", "₹", text)
    # Fonts that map ₹ to the glyph 'n': "up to n1,500" → "up to ₹1,500"
    # Only match standalone 'n' immediately before a number (not mid-word)
    text = _r.sub(r"(?<=[\s(])n(?=[\d,])", "₹", text)
    text = _r.sub(r" {3,}", "  ", text)
    return text


def extract_pdf(filepath: str) -> str:
    """
    Robust PDF extractor with 3 fallbacks:
      1. pdfplumber  (best for structured/text PDFs)
      2. PyMuPDF / fitz  (better for complex layouts, scanned docs)
      3. pdfminer  (pure-Python fallback, no binary deps)
    Returns the best non-empty result.
    """
    text_parts = []

    # ── Attempt 1: pdfplumber ─────────────────────────────────────────────────
    if HAS_PDF:
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t and t.strip():
                        text_parts.append(t.strip())
            if text_parts:
                combined = _clean_pdf_text("\n".join(text_parts).strip())
                if len(combined) > 100:      # meaningful content
                    print(f"    ✅ PDF via pdfplumber: {len(combined)} chars")
                    return combined
        except Exception as e:
            print(f"    ⚠️  pdfplumber failed: {e}")
        text_parts = []

    # ── Attempt 2: pypdfium2 ─────────────────────────────────────────────────
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(filepath)
        for page in doc:
            tb = page.get_textpage()
            t  = tb.get_text_range()
            if t and t.strip():
                text_parts.append(t.strip())
        if text_parts:
            combined = _clean_pdf_text("\n".join(text_parts).strip())
            if len(combined) > 100:
                print(f"    ✅ PDF via pypdfium2: {len(combined)} chars")
                return combined
    except ImportError:
        pass
    except Exception as e:
        print(f"    ⚠️  pypdfium2 failed: {e}")
    text_parts = []

    # ── Attempt 3: pdfminer ───────────────────────────────────────────────────
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(filepath)
        if text and len(text.strip()) > 100:
            text = _clean_pdf_text(text.strip())
            print(f"    ✅ PDF via pdfminer: {len(text)} chars")
            return text
    except ImportError:
        pass
    except Exception as e:
        print(f"    ⚠️  pdfminer failed: {e}")

    print(f"    ❌ All PDF extractors failed for: {filepath}")
    return ""



def extract_docx(filepath: str) -> str:
    """
    Extract text from DOCX, preserving heading structure.
    Word Heading styles (Heading 1/2/3) are output on their own line,
    which makes them detectable by parse_text_to_sections.
    Tables are extracted row-by-row.
    Falls back to raw XML zip extraction if python-docx fails.
    """
    if not HAS_DOCX:
        return ""
    try:
        doc = DocxDocument(filepath)
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if not t:
                continue
            style = (para.style.name or "").lower()
            # Word heading styles → output as standalone line (detected as heading)
            if "heading" in style or style == "title":
                parts.append("")          # blank line before heading
                parts.append(t)           # heading on its own line
                parts.append("")          # blank line after
            else:
                parts.append(t)

        # Tables
        for table in doc.tables:
            parts.append("")
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        result = "\n".join(parts)
        if result.strip():
            print(f"    ✅ DOCX extracted: {len(result)} chars")
        return result
    except Exception as e:
        print(f"    ❌ DOCX extraction error: {e}")
        try:
            import zipfile, xml.etree.ElementTree as ET
            with zipfile.ZipFile(filepath) as z:
                xml_content = z.read("word/document.xml")
            tree = ET.fromstring(xml_content)
            texts = [node.text for node in
                     tree.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                     if node.text]
            result = " ".join(texts)
            if result.strip():
                print(f"    ✅ DOCX fallback (raw XML): {len(result)} chars")
                return result
        except Exception as e2:
            print(f"    ❌ DOCX XML fallback also failed: {e2}")
        return ""


def extract_md(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    text = re.sub(r"#{1,6}\s*", "", text)
    text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"`{1,3}.*?`{1,3}", "", text, flags=re.DOTALL)
    return text


def extract_html(filepath: str) -> str:
    """Extract text from local HTML files."""
    if not HAS_BS4:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            return re.sub(r"<[^>]+>", " ", f.read())
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    return soup.get_text(separator="\n")


def extract_rtf(filepath: str) -> str:
    """Basic RTF extraction — strips RTF control words."""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    # Remove RTF control words and groups
    text = re.sub(r"\\[a-z]+\d*\s?", " ", text)
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"\\\*.*?;", "", text)
    return text.strip()


def _clean_web_text(soup_or_text, url: str = "") -> tuple[str, str]:
    """
    Given a BeautifulSoup object or raw HTML string, extract clean text.
    Returns (title, clean_text).
    """
    if isinstance(soup_or_text, str):
        soup = BeautifulSoup(soup_or_text, "html.parser")
    else:
        soup = soup_or_text

    title = soup.title.string.strip() if soup.title and soup.title.string else url

    for tag in soup(["script", "style", "nav", "footer", "header",
                     "aside", "form", "noscript", "iframe", "img"]):
        tag.decompose()

    main = (soup.find("main") or soup.find("article") or
            soup.find(id=re.compile(r"content|main", re.I)) or
            soup.body)
    text = main.get_text(separator="\n") if main else soup.get_text(separator="\n")
    lines = [l.strip() for l in text.splitlines() if l.strip() and len(l.strip()) > 3]
    return title, "\n".join(lines)


def extract_url(url: str, follow_links: bool = True) -> tuple[str, str]:
    """
    Fetch a URL, extract visible text, and optionally follow relevant sub-links
    on the same domain (depth-1 crawl) to get comprehensive content.
    Returns (scheme_name, full_text).
    """
    if not HAS_BS4:
        return ("", "")

    headers = {"User-Agent": "Mozilla/5.0 (IndicGovAdvisor/2.0)"}
    base_domain = urlparse(url).netloc

    try:
        r = requests.get(url, headers=headers, timeout=20)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        title, main_text = _clean_web_text(soup, url)

        all_text_parts = [main_text]

        # Depth-1: follow links on same domain that look like scheme content
        if follow_links:
            links = soup.find_all("a", href=True)
            visited = {url}
            scheme_keywords = re.compile(
                r"(scheme|yojana|benefit|eligib|apply|document|faq|about|overview)",
                re.IGNORECASE
            )
            followed = 0
            for a in links:
                if followed >= 5:  # max 5 sub-pages
                    break
                href = urljoin(url, a["href"])
                href_domain = urlparse(href).netloc
                if href in visited or href_domain != base_domain:
                    continue
                if not scheme_keywords.search(href) and not scheme_keywords.search(a.get_text()):
                    continue
                try:
                    sub_r = requests.get(href, headers=headers, timeout=15)
                    sub_r.raise_for_status()
                    _, sub_text = _clean_web_text(sub_r.text, href)
                    if sub_text and len(sub_text) > 200:
                        all_text_parts.append(f"\n--- Section from: {href} ---\n{sub_text}")
                        visited.add(href)
                        followed += 1
                        print(f"      ↳ Followed link: {href}")
                except Exception:
                    pass

        combined = "\n\n".join(all_text_parts)
        return (title, combined)

    except Exception as e:
        print(f"    ❌ Failed to fetch {url}: {e}")
        return ("", "")


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION PARSER
# ═══════════════════════════════════════════════════════════════════════════════

# ── Canonical section name map ───────────────────────────────────────────────
# Maps any heading variant → the standard section key used by RAG retrieval
_SECTION_CANON = [
    ("overview",              ["overview", "about", "introduction", "background", "summary", "scheme overview"]),
    ("benefits",              ["benefit", "advantages", "features", "what you get", "financial benefit",
                               "key benefits", "scheme benefits", "incentives"]),
    ("eligibility",           ["eligib", "who can apply", "who is eligible", "criteria", "qualification",
                               "target group", "beneficiar", "entitled"]),
    ("application_process",   ["application", "how to apply", "apply", "registration process",
                               "enrollment", "procedure", "steps to apply", "process"]),
    ("documents_required",    ["document", "required document", "supporting document", "papers needed",
                               "documents needed", "list of documents", "paperwork", "id proof"]),
    ("faqs",                  ["faq", "frequently asked", "common question", "q&a", "queries"]),
    ("website_url",           ["website", "web url", "official link", "portal", "online"]),
    ("financial_details",     ["amount", "fund", "financial", "budget", "cost", "rupee", "payment",
                               "disbursement", "grant", "subsidy", "loan amount"]),
    ("implementation",        ["implement", "nodal", "ministry", "department", "agency",
                               "launched by", "responsible body"]),
]

def _canonicalize_section(raw_heading: str) -> str:
    """Map any raw heading string to its canonical section key."""
    h = raw_heading.lower().strip()
    for canon_key, keywords in _SECTION_CANON:
        if any(kw in h for kw in keywords):
            return canon_key
    # Fallback: slugify the heading itself
    return re.sub(r"[^a-z0-9]+", "_", h).strip("_") or "general"


# Single-word headings common in PDFs and plain text
_SINGLE_WORD_HEADINGS = {
    "overview", "benefits", "eligibility", "introduction", "summary",
    "background", "documents", "faqs", "faq", "application", "process",
    "procedure", "features", "criteria", "objectives", "objective",
}

# Multi-word heading phrases used as section titles across government docs
_KNOWN_HEADING_PHRASES = {
    "key benefits", "scheme benefits", "main benefits", "financial benefits",
    "eligibility criteria", "who can apply", "who is eligible", "target beneficiaries",
    "application process", "how to apply", "steps to apply", "registration process",
    "documents required", "required documents", "documents needed", "list of documents",
    "frequently asked questions", "contact information", "important links",
    "financial details", "scheme overview", "official website", "website url",
    "implementation details", "nodal agency", "about the scheme",
}

def _is_heading(line: str) -> bool:
    """
    Detect section headings. Deliberately conservative to avoid treating
    content lines as headings.

    Accepts:
      1. Numbered:  "1. Overview:"  "4. Application Process:"
      2. Known single words: "Benefits", "Eligibility", "Overview"
      3. Known phrases: "Eligibility Criteria", "Documents Required"
      4. ALL CAPS 2+ words: "ELIGIBILITY CRITERIA"
      5. Any line ending with ":" that is 2–6 words and has no digits/urls
      6. Bold markdown: **Benefits**

    Rejects:
      - Long lines (>100 chars)
      - Lines containing digits, ₹, Rs., URLs, email
      - Any line that looks like a sentence (ends with . ! ?)
      - Lines with commas or semicolons (list content, not headings)
    """
    s = line.strip()
    if not s or len(s) > 100:
        return False
    sl = s.lower().rstrip(":")

    # 1. Numbered heading — reject if it ends with sentence punctuation
    #    (catches "3 Completion of antenatal check-ups." which is a list item, not a heading)
    if re.match(r"^\d+[.)\s]\s*.{3,60}$", s) and not re.search(r"[.!?]$", s):
        return True
    # 2. Known single-word headings
    if sl in _SINGLE_WORD_HEADINGS:
        return True
    # 3. Known multi-word heading phrases
    if sl in _KNOWN_HEADING_PHRASES:
        return True
    # 4. ALL CAPS, 2+ words, no digits
    if (re.match(r"^[A-Z][A-Z\s\-/()]{3,}$", s)
            and len(s.split()) >= 2
            and not re.search(r"\d", s)):
        return True
    # 5. Ends with colon, short (2–6 words), no digits/symbols/sentences
    words = s.split()
    if (s.endswith(":")
            and 2 <= len(words) <= 6
            and not re.search(r"[\d₹@]|Rs\.?|http", s)):
        return True
    # 6. Bold markdown **Heading** or __Heading__
    if re.match(r"^[*_]{2}.{3,60}[*_]{2}:?$", s):
        return True
    return False


def _extract_heading_text(line: str) -> str:
    """Strip numbering, colons, and bold markers from a heading line."""
    s = line.strip()
    s = re.sub(r"^\d+[.)\s]\s*", "", s)   # remove "1. " prefix
    s = re.sub(r"^[*_]{2}|[*_]{2}$", "", s)  # remove **bold** markers
    s = s.rstrip(":").strip()
    return s


def parse_text_to_sections(scheme_name: str, text: str) -> dict:
    """
    Split extracted text into named sections.
    Works for: structured .txt (numbered), Word docs (Title Case headings),
    PDFs (ALL CAPS or mixed headings), and unstructured text (paragraph fallback).
    """
    lines = text.splitlines()
    sections: dict = {}
    current_section = "general"
    current_lines: list = []

    for line in lines:
        if _is_heading(line):
            # Save accumulated content under current section
            if current_lines:
                body = "\n".join(current_lines).strip()
                if body and len(body) > 20:
                    sections[current_section] = (
                        sections.get(current_section, "") + "\n" + body
                    ).strip()
            # Start new section with canonical key
            heading_text  = _extract_heading_text(line)
            current_section = _canonicalize_section(heading_text)
            current_lines = []
        else:
            stripped = line.strip()
            if stripped and not stripped.startswith("---"):
                current_lines.append(line)

    # Flush last section
    if current_lines:
        body = "\n".join(current_lines).strip()
        if body and len(body) > 20:
            sections[current_section] = (
                sections.get(current_section, "") + "\n" + body
            ).strip()

    # ── Fallback: everything ended up in "general" (no headings detected) ──
    if len(sections) <= 1 and "general" in sections:
        paras = [p.strip() for p in re.split(r"\n{2,}", sections["general"]) if len(p.strip()) > 60]
        if len(paras) >= 3:
            # Try to assign semantic labels based on paragraph content
            labelled = {}
            for para in paras:
                pl = para.lower()
                if any(k in pl for k in ["eligib", "who can", "criteria", "qualify"]):
                    key = "eligibility"
                elif any(k in pl for k in ["benefit", "assist", "incentive", "cash", "₹", "rs."]):
                    key = "benefits"
                elif any(k in pl for k in ["apply", "application", "register", "enroll", "step"]):
                    key = "application_process"
                elif any(k in pl for k in ["document", "proof", "id card", "certificate"]):
                    key = "documents_required"
                else:
                    key = "overview"
                existing = labelled.get(key, "")
                labelled[key] = (existing + "\n\n" + para).strip() if existing else para
            if len(labelled) >= 2:
                sections = labelled

    return sections


def get_scheme_name_from_text(text: str, fallback: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped and not stripped.startswith("#"):
            return stripped
    return fallback


# ═══════════════════════════════════════════════════════════════════════════════
#  FILE DISPATCHER
# ═══════════════════════════════════════════════════════════════════════════════

EXTRACTORS = {
    ".txt":  extract_txt,
    ".md":   extract_md,
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".doc":  extract_docx,
    ".html": extract_html,
    ".htm":  extract_html,
    ".rtf":  extract_rtf,
}

SUPPORTED_EXTS = set(EXTRACTORS.keys())


def ingest_file(filepath: str, filename: str) -> tuple[str, dict]:
    ext = os.path.splitext(filename)[1].lower()
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        print(f"    ⏩ Skipping unsupported format: {filename}")
        return ("", {})

    raw = extractor(filepath)
    if not raw.strip():
        print(f"    ⚠️  Empty content from: {filename}")
        return ("", {})

    fallback_name = os.path.splitext(filename)[0].replace("_", " ").title()
    scheme_name   = get_scheme_name_from_text(raw, fallback_name)

    # Skip the standard 4-line metadata header (Ministry/Department/Launch/Type)
    # so these lines don't get misidentified as section headings
    lines = raw.splitlines()
    content_start = 0
    for i, line in enumerate(lines[:10]):
        l = line.strip().lower()
        if l.startswith(("ministry:", "department:", "launch date:", "type:", "---")):
            content_start = i + 1
    if content_start > 0:
        raw = lines[0] + "\n" + "\n".join(lines[content_start:])  # keep scheme name line

    sections      = parse_text_to_sections(scheme_name, raw)
    return (scheme_name, sections)


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN INGEST
# ═══════════════════════════════════════════════════════════════════════════════

def ingest_all():
    if not os.path.isdir(SCHEMES_DIR):
        print(f"⚠️  Schemes directory '{SCHEMES_DIR}' not found. Creating it.")
        os.makedirs(SCHEMES_DIR, exist_ok=True)
        return

    docs, metas, ids = [], [], []
    total_files  = 0
    total_chunks = 0

    # ── 1. Files ──────────────────────────────────────────────────────────────
    all_files = [
        f for f in os.listdir(SCHEMES_DIR)
        if os.path.splitext(f)[1].lower() in SUPPORTED_EXTS
        and f.lower() != "urls.txt"
    ]

    for filename in sorted(all_files):
        filepath = os.path.join(SCHEMES_DIR, filename)
        print(f"  📄 Ingesting file: {filename}")
        scheme_name, sections = ingest_file(filepath, filename)
        if not sections:
            continue
        total_files += 1

        for section_key, section_text in sections.items():
            if len(section_text.strip()) < 30:
                continue
            chunk_id = f"{filename}::{section_key}"
            docs.append(section_text)
            metas.append({
                "scheme_name": scheme_name,
                "section":     section_key,
                "source_file": filename,
            })
            ids.append(chunk_id)
            total_chunks += 1

    # ── 2. URLs from urls.txt ─────────────────────────────────────────────────
    urls_file = os.path.join(SCHEMES_DIR, "urls.txt")
    if os.path.exists(urls_file) and HAS_BS4:
        with open(urls_file, "r", encoding="utf-8") as f:
            urls = [l.strip() for l in f if l.strip() and not l.startswith("#")]

        for url in urls:
            print(f"  🌐 Crawling URL: {url}")
            title, text = extract_url(url, follow_links=True)
            if not text:
                continue
            sections = parse_text_to_sections(title, text)
            total_files += 1

            for section_key, section_text in sections.items():
                if len(section_text.strip()) < 30:
                    continue
                safe_url = re.sub(r"[^\w]", "_", url)[:80]
                chunk_id = f"{safe_url}::{section_key}"
                docs.append(section_text)
                metas.append({
                    "scheme_name": title,
                    "section":     section_key,
                    "source_file": url,
                })
                ids.append(chunk_id)
                total_chunks += 1

    elif os.path.exists(urls_file) and not HAS_BS4:
        print("⚠️  urls.txt found but beautifulsoup4 not installed — skipping URL crawl.")

    if not docs:
        print(f"⚠️  No content found. Add files to '{SCHEMES_DIR}/'")
        return

    collection.add(documents=docs, metadatas=metas, ids=ids)
    print(f"\n✅ Ingested {total_files} sources → {total_chunks} chunks into ChromaDB")
    print(f"   Database saved at: {CHROMA_DIR}/")


if __name__ == "__main__":
    print("🔄 Starting ingestion...\n")
    ingest_all()